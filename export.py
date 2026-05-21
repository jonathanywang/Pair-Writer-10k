"""Markdown and DOCX export for the memo."""


from __future__ import annotations
import ssl_patch  # ← line 1 of every file

import io
import re

from docx import Document
from docx.shared import Pt

from edits import Memo


def memo_to_markdown(memo: Memo, title: str = "Investment Memo") -> str:
    """Render the memo as Markdown. Citations stay as [1A·¶3]."""
    lines = [f"# {title}", ""]
    for p in memo.paragraphs:
        lines.append(p.text)
        lines.append("")
    return "\n".join(lines)


def memo_to_docx(memo: Memo, title: str = "Investment Memo",
                 filing_name: str = "10-K") -> bytes:
    """Render the memo as a DOCX file. Citations rendered as inline parentheticals."""
    doc = Document()

    # Title
    title_para = doc.add_heading(title, level=1)

    # Subtitle with filing name
    sub = doc.add_paragraph()
    run = sub.add_run(f"Source filing: {filing_name}")
    run.italic = True
    run.font.size = Pt(10)

    doc.add_paragraph()  # spacer

    # Body paragraphs
    citation_pattern = re.compile(r"\[([0-9A-Z]+)·([¶pPtT])(\d+)\]")

    for p in memo.paragraphs:
        # Replace citations with inline parentheticals
        text = citation_pattern.sub(
            lambda m: f" ({filing_name}, Item {m.group(1)}, {'¶' if m.group(2) in ('¶', 'p', 'P') else 'Table '}{m.group(3)})",
            p.text,
        )
        doc.add_paragraph(text)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
